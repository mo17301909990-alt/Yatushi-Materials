#!/usr/bin/env python3
"""检查 docx 文件和更多表头细节"""

import io, sys

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

# 检查 docx
try:
    import docx

    doc = docx.Document("G:/雅图仕/download/LEO纸品/LEO-各品牌纸类通用产品指引书-202602012.docx")
    print(f"=== DOCX 文件 ===")
    print(f"  段落数: {len(doc.paragraphs)}")
    print(f"  表格数: {len(doc.tables)}")
    print(f"  前20段落:")
    for i, p in enumerate(doc.paragraphs[:20]):
        if p.text.strip():
            print(f"    [{i}] {p.text.strip()[:100]}")
    print(f"  前3表格行列数:")
    for i, t in enumerate(doc.tables[:3]):
        print(f"    表格{i}: {len(t.rows)}行 x {len(t.columns)}列")
except ImportError:
    print("python-docx not available")
except Exception as e:
    print(f"docx error: {e}")

# 更详细检查 LEO 格式的"界面"(surface) 子表头结构
base = "G:/雅图仕/download/LEO纸品"
import openpyxl
from openpyxl.utils import get_column_letter

files = [
    ("书板贴纸", f"{base}/书板贴纸-应用指引书-20251201 (2).xlsx"),
    ("平面产品贴纸", f"{base}/平面产品贴纸-应用指引书-20250701 (7).xlsx"),
    ("非平面产品贴纸", f"{base}/非平面产品贴纸-应用指引书-20250701 (7).xlsx"),
]

def sanitize(v):
    if v is None:
        return ""
    return str(v).strip().replace("\n", "\\n")[:60]

for label, path in files:
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb[wb.sheetnames[0]]
    print(f"\n{'='*80}")
    print(f"=== {label} — 全表头行结构 ===")
    print(f"{'='*80}")

    # 打印 15-25 行的所有非空列（按列号排序）
    for r in range(15, min(26, ws.max_row + 1)):
        cells = []
        for c in range(1, min(ws.max_column + 1, 160)):
            s = sanitize(ws.cell(r, c).value)
            if s:
                cells.append(f"{get_column_letter(c)}({c})={s}")
        if cells:
            print(f"\n  Row {r}:")
            for item in cells[:25]:  # 太多列时只显前25
                print(f"    {item}")
            if len(cells) > 25:
                print(f"    ... 还有 {len(cells)-25} 个非空列")

    # 检查 step_row (find_step_header_row 能找到的行)
    # step_keywords check: rows 15-30, cols 40-120, >=20 non-empty
    print(f"\n  --- 模拟 find_step_header_row ---")
    step_keywords = {"單面印刷", "雙面印刷", "連線印刷", "離線印刷", "油性油",
        "水性油", "UV油", "普通油", "普通燙金", "鐳射燙金",
        "立體燙金", "擊凸", "擊凹", "過膠", "絲印", "植毛",
        "機啤", "手啤", "壓紋", "粘貼", "車線"}
    for r in range(15, min(31, ws.max_row + 1)):
        count = 0
        hits = []
        for c in range(40, min(ws.max_column + 1, 121)):
            s = sanitize(ws.cell(r, c).value)
            if s:
                count += 1
                if any(kw in s for kw in step_keywords) or (len(s) >= 2 and len(s) <= 20):
                    hits.append(f"{get_column_letter(c)}='{s}'")
        if count > 5:
            print(f"    Row {r}: {count} 个非空列, {len(hits)} 个关键词命中")
            if hits:
                for h in hits[:8]:
                    print(f"      {h}")

    wb.close()
